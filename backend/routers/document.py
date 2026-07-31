import io
import logging

from docx import Document
from docx.shared import Pt
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from services.humanize_service import humanize as _humanize
from services.paraphrase_service import paraphrase as _paraphrase

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["document"])


class DocProcessResponse(BaseModel):
    original_text: str
    processed_text: str
    mode: str
    filename: str


class DocExtractResponse(BaseModel):
    text: str
    filename: str
    character_count: int


class DocDownloadRequest(BaseModel):
    processed_text: str
    filename: str


def _extract_document_text(filename: str, content: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "docx":
        try:
            doc = Document(io.BytesIO(content))
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Could not read .docx file: {exc}")
        text = "\n\n".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
    elif suffix == "txt":
        try:
            text = content.decode("utf-8-sig").strip()
        except UnicodeDecodeError as exc:
            raise HTTPException(status_code=400, detail="Text files must use UTF-8 encoding") from exc
    else:
        raise HTTPException(status_code=400, detail="Upload a .docx or .txt document")

    if not text:
        raise HTTPException(status_code=400, detail="Document contains no readable text")
    if len(text) > 50000:
        raise HTTPException(
            status_code=400,
            detail="Document too large. Maximum 50,000 characters supported.",
        )
    return text


@router.post(
    "/documents/extract",
    response_model=DocExtractResponse,
    summary="Extract editable text from a document",
)
async def extract_document(file: UploadFile = File(...)) -> DocExtractResponse:
    """Load a document into the unified writing workspace without processing it."""
    filename = file.filename or "document.docx"
    text = _extract_document_text(filename, await file.read())
    return DocExtractResponse(text=text, filename=filename, character_count=len(text))


@router.post(
    "/upload-doc",
    response_model=DocProcessResponse,
    summary="Upload a Word doc and humanize or paraphrase it",
)
async def upload_doc(
    file: UploadFile = File(...),
    mode: str = Form("humanize"),
):
    """
    Upload a .docx file and process it through humanize or paraphrase.

    Returns JSON with original and processed text for side-by-side comparison.
    Use `/api/upload-doc/download` to get the result as a .docx file.
    """
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported by this legacy endpoint")

    if mode not in ("humanize", "paraphrase"):
        raise HTTPException(status_code=400, detail="Mode must be 'humanize' or 'paraphrase'")

    full_text = _extract_document_text(file.filename, await file.read())

    try:
        if mode == "humanize":
            result = _humanize(full_text)
            processed_text = result["humanized"]
        else:
            processed_text, _ = _paraphrase(full_text, intensity=3)
    except Exception as exc:
        logger.exception("Document processing failed")
        raise HTTPException(status_code=500, detail=f"Processing failed: {exc}")

    return DocProcessResponse(
        original_text=full_text,
        processed_text=processed_text,
        mode=mode,
        filename=file.filename or "document.docx",
    )


@router.post("/upload-doc/download", summary="Download processed text as .docx")
def download_doc(request: DocDownloadRequest):
    """Convert processed text back into a downloadable .docx file."""
    out_doc = Document()
    style = out_doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    for paragraph in request.processed_text.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            out_doc.add_paragraph(paragraph)

    buf = io.BytesIO()
    out_doc.save(buf)
    buf.seek(0)

    base_name = request.filename.rsplit(".", 1)[0]
    out_name = f"{base_name}_processed.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )
